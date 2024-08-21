from docx import Document
import os
#doc = Document(r'C:\Users\jimee\Desktop\ttyy\北京工业厂房施工组织设计.docx')

def get_docx_info(file_path):
    doc = Document(file_path)
    paragraph_count = len(doc.paragraphs)
    file_size = os.path.getsize(file_path) / 1024

    print(f"文件大小：{file_size:.2f} KB")
    print(f"段落数量（约等于页数）：{paragraph_count}")

if __name__ == '__main__':
    file_path = r'C:\Users\jimee\Desktop\ttyy\甘肃省屋面施工方案.docx'
    get_docx_info(file_path)